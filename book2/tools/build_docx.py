# -*- coding: utf-8 -*-
"""Minimal, format-aware Markdown -> DOCX for the rhythm-game textbook sample.
Handles: #/##/### headings, paragraphs w/ **bold** and `code`, pipe tables,
fenced code (prompt boxes), admonitions (!!! kind "title"), blockquotes (FIG
captions), images, lists, workbook lines. Tailored to the book's own syntax."""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
BOOK = os.path.dirname(HERE)                 # book2/
LESS = os.path.join(BOOK, "lessons")
FILES = sorted(x for x in os.listdir(LESS) if x.endswith(".md"))
OUT = os.path.join(BOOK, "리듬게임교재_전체_v3.docx")
KOFONT = "맑은 고딕"

ADMON = {  # kind -> (title color, shade fill)
    "note":     ("1F4E79", "DDEBF7"),
    "tip":      ("2E7D32", "E2F0D9"),
    "warning":  ("9C6500", "FFF2CC"),
    "success":  ("1E7145", "E2EFDA"),
    "question": ("6A1B9A", "EAE0F3"),
    "info":     ("444444", "EDEDED"),
    "abstract": ("0F6B6B", "D6EFEF"),
    "example":  ("8A5A00", "FCE9D6"),
    "quote":    ("7A5C00", "FFF7D6"),  # prompt box — distinct
}
KIND_ICON = {"quote": "🗣️", "note": "📓", "tip": "💡", "warning": "⚠️",
             "success": "✅", "question": "🤔", "info": "🔎", "abstract": "📌",
             "example": "✏️"}

def set_font(run, name=KOFONT, size=None, bold=None, color=None, mono=False):
    fn = "Consolas" if mono else name
    run.font.name = fn
    r = run._element.rPr
    if r is None:
        r = run._element.get_or_add_rPr()
    rf = r.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); r.append(rf)
    rf.set(qn('w:ascii'), fn); rf.set(qn('w:hAnsi'), fn)
    rf.set(qn('w:eastAsia'), name)  # keep Korean glyphs in KOFONT even in code
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)

def add_inline(p, text, size=None, bold_all=False, color=None, mono=False):
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1', text)  # [t](u)->t
    text = text.replace('✍️', '✍').replace('☐', '□')
    for part in re.split(r'(\*\*.+?\*\*|`[^`]+`)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2]); set_font(run, size=size, bold=True, color=color)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1]); set_font(run, size=size, bold=bold_all, color="B03A00", mono=True)
        else:
            run = p.add_run(part); set_font(run, size=size, bold=bold_all, color=color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def borders(cell, color="BFBFBF", sz="8"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz); e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    tcPr.append(b)

def tokenize(lines):
    toks, i, n = [], 0, len(lines)
    while i < n:
        ln = lines[i]
        s = ln.strip()
        if s.startswith('<!--'):
            i += 1; continue
        if s == '---':
            toks.append(('hr', None)); i += 1; continue
        if s == '':
            i += 1; continue
        if ln.startswith('### '):
            toks.append(('h3', s[4:])); i += 1; continue
        if ln.startswith('## '):
            toks.append(('h2', s[3:])); i += 1; continue
        if ln.startswith('# '):
            toks.append(('h1', s[2:])); i += 1; continue
        m = re.match(r'!!!\s+(\w+)\s+"(.*)"', s)
        if m:
            kind, title = m.group(1), m.group(2)
            i += 1; body = []
            while i < n and (lines[i].strip() == '' or lines[i].startswith('    ')):
                body.append(lines[i][4:] if lines[i].startswith('    ') else '')
                i += 1
            while body and body[-1].strip() == '':
                body.pop()
            toks.append(('admon', (kind, title, body))); continue
        if s.startswith('```'):
            i += 1; code = []
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1
            toks.append(('code', code)); continue
        if s.startswith('|'):
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                rows.append(lines[i].strip()); i += 1
            toks.append(('table', rows)); continue
        if s.startswith('>'):
            q = []
            while i < n and lines[i].strip().startswith('>'):
                q.append(lines[i].strip()[1:].strip()); i += 1
            toks.append(('quote', q)); continue
        mi = re.match(r'!\[(.*?)\]\((.*?)\)', s)
        if mi:
            toks.append(('img', (mi.group(1), mi.group(2)))); i += 1; continue
        if re.match(r'(-|\*|\d+\.)\s+', s):
            items = []
            while i < n and re.match(r'(-|\*|\d+\.)\s+', lines[i].strip()):
                items.append(re.sub(r'^(-|\*|\d+\.)\s+', '', lines[i].strip())); i += 1
            toks.append(('list', items)); continue
        # paragraph (gather following plain lines)
        para = [s]; i += 1
        while i < n:
            t = lines[i].strip()
            if t == '' or t.startswith(('#', '!!!', '```', '|', '>', '-', '*')) or t.startswith('<!--') or re.match(r'\d+\.\s', t):
                break
            para.append(t); i += 1
        toks.append(('para', ' '.join(para)))
    return toks

def render(container, toks, doc):
    for kind, data in toks:
        if kind == 'h1':
            h = container.add_paragraph(); h.space_before = Pt(6)
            r = h.add_run(data); set_font(r, size=20, bold=True, color="1F3864")
        elif kind == 'h2':
            h = container.add_paragraph()
            r = h.add_run(data); set_font(r, size=15, bold=True, color="2E5496")
            pPr = h._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bo = OxmlElement('w:bottom')
            bo.set(qn('w:val'), 'single'); bo.set(qn('w:sz'), '6'); bo.set(qn('w:space'), '2'); bo.set(qn('w:color'), 'C7D3E8')
            pbdr.append(bo); pPr.append(pbdr)
        elif kind == 'h3':
            h = container.add_paragraph()
            r = h.add_run(data); set_font(r, size=13, bold=True, color="374A6B")
        elif kind == 'para':
            p = container.add_paragraph(); add_inline(p, data, size=11)
        elif kind == 'list':
            for it in data:
                p = container.add_paragraph(style=None); p.paragraph_format.left_indent = Inches(0.3)
                b = p.add_run("•  "); set_font(b, size=11)
                add_inline(p, it, size=11)
        elif kind == 'code':
            tbl = container.add_table(rows=1, cols=1); tbl.autofit = True
            c = tbl.rows[0].cells[0]; shade(c, "F2F2F2"); borders(c, "D9D9D9")
            first = True
            for cl in (data or ['']):
                p = c.paragraphs[0] if first else c.add_paragraph(); first = False
                r = p.add_run(cl); set_font(r, size=10.5, mono=True, color="333333")
        elif kind == 'quote':
            for q in data:
                if not q:
                    continue
                p = container.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
                add_inline(p, q, size=10.5, color="555555")
        elif kind == 'img':
            alt, path = data
            ap = os.path.normpath(os.path.join(LESS, path))
            if os.path.exists(ap):
                pic = container.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic.add_run().add_picture(ap, width=Inches(6.0))
            else:
                p = container.add_paragraph(); r = p.add_run(f"[이미지 준비중: {alt}]"); set_font(r, size=10, color="999999")
        elif kind == 'table':
            render_table(container, data)
        elif kind == 'hr':
            container.add_paragraph()
        elif kind == 'admon':
            kd, title, body = data
            col, fill = ADMON.get(kd, ("444444", "EDEDED"))
            tbl = container.add_table(rows=1, cols=1); tbl.autofit = True
            c = tbl.rows[0].cells[0]; shade(c, fill); borders(c, col)
            tp = c.paragraphs[0]
            icon = KIND_ICON.get(kd, "")
            tr = tp.add_run((icon + "  " if icon and not title.startswith(tuple("🗣️💡⚠️✅🤔🔎📌✏️📓")) else "") + title)
            set_font(tr, size=11.5, bold=True, color=col)
            render(c, tokenize(body), doc)
        else:
            pass

def render_table(container, rows):
    cells = [[x.strip() for x in r.strip().strip('|').split('|')] for r in rows]
    cells = [c for c in cells if not all(re.match(r':?-+:?$', x or '-') for x in c)]
    if not cells:
        return
    ncol = max(len(r) for r in cells)
    t = container.add_table(rows=len(cells), cols=ncol)
    t.style = 'Table Grid'
    for ri, row in enumerate(cells):
        for ci in range(ncol):
            cell = t.rows[ri].cells[ci]
            cell.paragraphs[0].clear() if cell.paragraphs[0].runs else None
            txt = row[ci] if ci < len(row) else ''
            p = cell.paragraphs[0]
            add_inline(p, txt, size=10, bold_all=(ri == 0))
            if ri == 0:
                shade(cell, "1F3864")
                for run in p.runs:
                    run.font.color.rgb = RGBColor.from_string("FFFFFF")

# ---- build ----
doc = Document()
st = doc.styles['Normal']; st.font.name = KOFONT; st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn('w:eastAsia'), KOFONT)
sec = doc.sections[0]
sec.top_margin = Inches(0.9); sec.bottom_margin = Inches(0.9)
sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)

# cover line
cv = doc.add_paragraph(); cv.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cv.add_run("《말로 만드는 내 키링 리듬게임》 — 집필 샘플"); set_font(r, size=22, bold=True, color="1F3864")
cv2 = doc.add_paragraph(); cv2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cv2.add_run("서문 · 1차시 · 2차시  (현재까지 작성분 · 초안)"); set_font(r, size=12, color="666666")
doc.add_paragraph()

for idx, fn in enumerate(FILES):
    with open(os.path.join(LESS, fn), encoding='utf-8') as f:
        lines = f.read().split('\n')
    if idx > 0:
        doc.add_page_break()
    render(doc, tokenize(lines), doc)

doc.save(OUT)
print("SAVED", OUT)
print("SIZE", os.path.getsize(OUT), "bytes")
